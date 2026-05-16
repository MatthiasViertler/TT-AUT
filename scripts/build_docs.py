"""
Converts docs/kapitalertrag_documentation.md → docs/kapitalertrag_documentation.docx.

Run at major version bumps to keep the human-readable Word doc current.
Usage:
    python scripts/build_docs.py
"""
import re
import subprocess
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
except ImportError:
    raise SystemExit("python-docx not installed. Run: pip install python-docx")

SRC = Path(__file__).parent.parent / "docs" / "kapitalertrag_documentation.md"
DST = Path(__file__).parent.parent / "docs" / "kapitalertrag_documentation.docx"

# Colour palette
NAVY    = RGBColor(0x1F, 0x38, 0x64)  # cover title / Heading 1
BLUE    = RGBColor(0x2E, 0x75, 0xB6)  # cover subtitle
GRAY    = RGBColor(0x40, 0x40, 0x40)  # cover tagline / meta
TH_BG   = "2E74B5"                    # table header background
TH_FG   = RGBColor(0xFF, 0xFF, 0xFF)  # header text white
TR_ALT  = "EBF3FB"                    # alternating data-row tint
CODE_BG = "F2F2F2"                    # code block background
CODE_FG = RGBColor(0x24, 0x24, 0x24)  # code text
MONO_FG = RGBColor(0xC0, 0x20, 0x40)  # inline `code` colour
LINK_FG = RGBColor(0x2E, 0x74, 0xB5)  # link colour


def _git_version() -> str:
    try:
        return subprocess.check_output(
            ["git", "describe", "--tags", "--abbrev=0"],
            text=True, stderr=subprocess.DEVNULL
        ).strip()
    except Exception:
        return "see git log"


def _shade_para(p, hex_fill: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def _shade_cell(cell, hex_fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_align(para, alignment) -> None:
    para.alignment = alignment


_INLINE = re.compile(r'\*\*(.+?)\*\*|`([^`]+)`|\[([^\]]+)\]\([^)]*\)')


def _add_inline(para, text: str, *, cell_size: int | None = None) -> None:
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            if cell_size:
                run.font.size = Pt(cell_size)
        btext, ctext, ltext = m.group(1), m.group(2), m.group(3)
        if btext is not None:
            run = para.add_run(btext)
            run.bold = True
            if cell_size:
                run.font.size = Pt(cell_size)
        elif ctext is not None:
            run = para.add_run(ctext)
            run.font.name = "Courier New"
            run.font.size = Pt((cell_size or 11) - 1)
            run.font.color.rgb = MONO_FG
        elif ltext is not None:
            run = para.add_run(ltext)
            run.font.color.rgb = LINK_FG
            if cell_size:
                run.font.size = Pt(cell_size)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        if cell_size:
            run.font.size = Pt(cell_size)


def _cover_para(doc, text: str, size: float, color: RGBColor,
                bold: bool = False, space_after: float = 6.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _render_toc(doc: Document, entries: list[str]) -> None:
    """Emit a styled Table of Contents page followed by a page break."""
    # Heading
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("Table of Contents")

    doc.add_paragraph()  # small gap

    # Each entry — strip markdown link syntax, keep text only
    link_re = re.compile(r'\[([^\]]+)\]\([^)]*\)')
    for entry in entries:
        text = link_re.sub(r'\1', entry)   # [text](url) → text
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(text)
        run.font.size = Pt(11)

    # Page break after ToC
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _render_cover_page(doc: Document, elements: list[tuple[str, str]]) -> None:
    """Emit a styled cover page followed by a page break."""
    version = _git_version()

    # --- collect the three logical cover lines from elements ---
    app_name   = "Kapitalertrag"
    app_sub    = "Austrian Capital Gains Tax Calculator"   # blue subtitle
    app_tag    = "Complete Setup & Usage Guide"            # gray tagline
    meta_lines: list[str] = []

    for kind, text in elements:
        if kind == "h1":
            parts = text.split(" — ", 1)
            app_name = parts[0].strip()
            if len(parts) > 1:
                app_tag = parts[1].strip()
        elif kind == "bold":
            app_sub = text
        elif kind == "plain":
            # Replace generic version placeholder with real git tag
            line = re.sub(r"Version:.*", f"Version: {version}", text)
            meta_lines.append(line)

    # Top spacer (matches old doc's ~2000 twip after-space)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(100)

    # Main title  — 32pt bold navy, centered
    _cover_para(doc, app_name, 32, NAVY, bold=True, space_after=10)

    # Blue subtitle line
    _cover_para(doc, app_sub, 14, BLUE, space_after=6)

    # Gray tagline
    _cover_para(doc, app_tag, 12, GRAY, space_after=30)

    # Thin separator rule effect via spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

    # Metadata lines (version, jurisdiction)
    for line in meta_lines:
        _cover_para(doc, line, 10, GRAY, space_after=4)

    # Page break — end of cover page
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _render_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=0, cols=ncols)
    tbl.style = "Table Grid"

    tr = tbl.add_row()
    for i, txt in enumerate(rows[0]):
        cell = tr.cells[i]
        _shade_cell(cell, TH_BG)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        _add_inline(p, txt, cell_size=9)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = TH_FG

    for ridx, row_data in enumerate(rows[1:], start=1):
        tr = tbl.add_row()
        if ridx % 2 == 0:
            for cell in tr.cells:
                _shade_cell(cell, TR_ALT)
        for i in range(ncols):
            txt = row_data[i] if i < len(row_data) else ""
            cell = tr.cells[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _add_inline(p, txt, cell_size=9)

    doc.add_paragraph()


def build() -> None:
    md = SRC.read_text(encoding="utf-8")
    doc = Document()

    for sec in doc.sections:
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)
        sec.top_margin    = Inches(0.9)
        sec.bottom_margin = Inches(0.9)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Parse state
    state      = "cover"    # cover → toc_collect → normal
    cover_els: list[tuple[str, str]] = []
    toc_entries: list[str] = []

    in_code    = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_table   = False

    def flush_code() -> None:
        if not code_lines:
            return
        p = doc.add_paragraph()
        _shade_para(p, CODE_BG)
        p.paragraph_format.left_indent  = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = CODE_FG
        code_lines.clear()

    def flush_table() -> None:
        if table_rows:
            _render_table(doc, list(table_rows))
            table_rows.clear()

    def is_hr(line: str) -> bool:
        return bool(re.match(r"^---+$", line.strip()))

    for line in md.splitlines():
        # ── Cover phase: collect elements, render on first HR ────────────
        if state == "cover":
            if is_hr(line):
                _render_cover_page(doc, cover_els)
                state = "toc_collect"
            elif line.startswith("# "):
                cover_els.append(("h1", line[2:]))
            elif re.match(r"^\*\*(.+)\*\*$", line.strip()):
                cover_els.append(("bold", line.strip().strip("*")))
            elif line.strip():
                cover_els.append(("plain", line.strip()))
            continue

        # ── ToC-collect phase: gather entries until next HR ─────────────
        if state == "toc_collect":
            if is_hr(line):
                _render_toc(doc, toc_entries)
                state = "normal"
            elif re.match(r"^\d+\.\s+\[", line):
                # numbered ToC entry: strip leading "N. "
                toc_entries.append(re.sub(r"^\d+\.\s+", "", line.strip()))
            # skip heading and empty lines within the ToC block
            continue

        # ── Normal rendering ─────────────────────────────────────────────
        if line.startswith("```"):
            if in_table:
                flush_table()
                in_table = False
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c.replace(" ", "")) for c in cells if c):
                in_table = True
                continue
            in_table = True
            table_rows.append(cells)
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        if is_hr(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 4")
            _add_inline(p, line[5:])
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            _add_inline(p, line[4:])
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            _add_inline(p, line[3:])
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            _add_inline(p, line[2:])
            continue

        m_ol = re.match(r"^(\d+)\.\s+(.+)$", line)
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, m_ol.group(2))
            continue

        if re.match(r"^  [-*] ", line):
            p = doc.add_paragraph(style="List Bullet 2")
            _add_inline(p, line[4:])
            continue

        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:])
            continue

        if not line.strip():
            continue

        p = doc.add_paragraph()
        _add_inline(p, line)

    flush_code()
    flush_table()

    doc.save(DST)
    print(f"Written: {DST}")


if __name__ == "__main__":
    build()
